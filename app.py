from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt, RGBColor
from docx2pdf import convert
import yagmail

app = Flask(__name__)


@app.route('/')
@app.route('/home')
def home():
    return render_template('index.html')


@app.route('/certificate')
def certificate():
    return render_template('cert.html')


@app.route('/report-gen')
def report():
    return render_template('report.html')


@app.route('/generate-certificate', methods=['POST'])
def generate_certificate():
    name = request.form['name']
    course = request.form['course']
    date = request.form['date']
    certificate_type = request.form['certificate_type']
    organization_name = request.form['organization_name']
    organizer_name = request.form['organizer_name']
    organizer_designation = request.form['organizer_designation']
    recipient_email = request.form['recipient_email']
    # Load template
    doc = Document('certificate_template.docx')

    # Replace fields in the template
    for p in doc.paragraphs:
        if 'Name' in p.text:
            p.text = p.text.replace('Name', name)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(32)  # Change the font size to 24
            p.runs[0].font.color.rgb = RGBColor(
                0x00, 0x00, 0x80)  # Set the font color to blue
        if 'CERTIFICATE_OF' in p.text:
            p.text = p.text.replace('CERTIFICATE_OF', certificate_type)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(35)
            p.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x80)
        if 'Course' in p.text:
            p.text = p.text.replace('Course', course)
            p.runs[0].font.size = Pt(26)
        if 'Date' in p.text:
            p.text = p.text.replace('Date', date)
            p.runs[0].font.size = Pt(26)
        if 'Organization_name' in p.text:
            p.text = p.text.replace('Organization_name', organization_name)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(28)
        if 'Organizer_name' in p.text:
            p.text = p.text.replace('Organizer_name', organizer_name)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(24)
        if 'Organizer_designation' in p.text:
            p.text = p.text.replace(
                'Organizer_designation', organizer_designation)
            p.runs[0].font.size = Pt(24)

    # Save the Word document
    doc.save('certificate.docx')
    convert("certificate.docx")
    certificate_path = 'certificate.pdf'
    # Send the email with the certificate
    is_sent = send_email_with_certificate(recipient_email, certificate_path)

    print(is_sent)

    if is_sent:
        return render_template('success.html')
    else:
        return render_template('error.html')


def send_email_with_certificate(recipient_email, certificate_path):
    # Configure yagmail
    sender_email = 'certgenproject@gmail.com'  # Replace with your email address
    # Replace with your email password or app password
    sender_password = 'bonqeueunpvqsvzz'
    subject = 'Certificate of Participation'

    try:
        is_sent = True
        # Create a yagmail object
        yag = yagmail.SMTP(sender_email, sender_password)

        # Send the email
        yag.send(
            to=recipient_email,
            subject=subject,
            contents='Please find the attached certificate.',
            attachments=[certificate_path]
        )
    except Exception as e:
        is_sent = False

    return is_sent


@app.route('/generate-report', methods=['POST'])
def generate_report():
    date = request.form['date']
    academicyear = request.form['academicyear']
    semester = request.form['semester']
    nameofevent = request.form['nameofevent']
    dateandtime = request.form['dateandtime']
    eventvenue = request.form['eventvenue']
    organizedby = request.form['organizedby']
    targetaudience = request.form['targetaudience']
    resourceperson = request.form['resourceperson']
    eventcontents = request.form['eventcontents']
    detailsofevent = request.form['detailsofevent']

    # Load template
    doc = Document('reportgentemplate.docx')

    # Replace fields in the template
    for p in doc.paragraphs:
        if 't1' in p.text:
            p.text = p.text.replace('t1', date)
        if 't2' in p.text:
            p.text = p.text.replace('t2', academicyear)
            p.runs[0].bold = True
        if 't3' in p.text:
            p.text = p.text.replace('t3', semester)
            p.runs[0].bold = True
        if 't4' in p.text:
            p.text = p.text.replace('t4', nameofevent)
        if 'tf1' in p.text:
            p.text = p.text.replace('tf1', eventcontents)
        if 'r1' in p.text:
            p.text = p.text.replace('r1', detailsofevent)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    if 't5' in cell_text:
                        cell.text = cell_text.replace('t5', dateandtime)
                    if 't6' in cell_text:
                        cell.text = cell_text.replace('t6', eventvenue)
                    if 't7' in cell_text:
                        cell.text = cell_text.replace('t7', organizedby)
                    if 't8' in cell_text:
                        cell.text = cell_text.replace('t8', targetaudience)
                    if 't9' in cell_text:
                        cell.text = cell_text.replace('t9', resourceperson)

    # Save the Word document
    doc.save('report.docx')

    if (send_file('report.docx', as_attachment=True)):
        # Send the file as a response
        return render_template('success.html')
    else:
        return render_template('error.html')


if __name__ == '__main__':

    app.run(debug=True, port=5502)
